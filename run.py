import asyncio
import logging
import random

from aiogram import Bot, Dispatcher
from aiogram.fsm.storage.memory import MemoryStorage
import os
from datetime import datetime
from aiogram import Bot
from aiogram import types, F
from aiogram.filters import CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.types import Message, FSInputFile
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.utils.keyboard import ReplyKeyboardBuilder, InlineKeyboardBuilder
from docx import Document
from docx.shared import Mm
from docx.shared import Pt

TOKEN = ''
bot = Bot(token=TOKEN)
dp = Dispatcher(storage=MemoryStorage())



edit_defects = {
    1: {'qst': 'Место расположения'},
    2: {'qst': 'Описание дефекта'},
    3: {'qst': 'Фото'},
    4: {'qst': ' Вероятная причина возникновения '},
    5: {'qst': 'Рекомендации'},
    6: {'qst': 'Завершить редактирование ↩️'},
}

new_defects = {
    1: {'nd': 'Место расположения'},
    2: {'nd': 'Описание дефекта'},
    3: {'nd': 'Фото'},
    4: {'nd': ' Вероятная причина возникновения '},
    5: {'nd': 'Рекомендации'},
    6: {'nd': '⬇ Сдедующий дефект ⬇'},
}

menu = {
    1: {'ans': 'Загрузить файл 📖'},
    2: {'ans': 'Сохранить файл 📝'},
    3: {'ans': 'Создать новый проект 🆕'}
}

privet = ['Доброго времени суток',
          'Будь как дома путник',
          'Welcome',
          'Добро пожаловать',
          'Рады вас видеть',
          'Hello',
          'Давно не видели вас в уличных гонках',
          'С возвращением',
          'Приятного пользования',
          'Рады приветствовать вас',
          'Мы вас ждали',
          'Бот готов к эксплуатации',
          'А я всё думал, когда же вы появитесь',
          'Наша радость от вашего визита не знает границ!']

stickers = ['CAACAgIAAxkBAAENr9tnnh01FnARp-4cYGG5RYdDVHYM4AACOAsAAk7kmUsysUfS2U-M0DYE',
           'CAACAgIAAxkBAAENwMBnqPQ5Io12GgsOx6SqIRDntKXhtgAC7goAAp2hYUuW93rCAubsUDYE',
           'CAACAgIAAxkBAAENvGpnpkIuQLOdaHk3JDotVtzmtMyXIgACsw4AAln8mUtNMf9WVqKCDjYE',
           'CAACAgIAAxkBAAENu59npeAfjF3pzKmPTgdY1RXBEhrgKAAC-woAAm80oUssauxdTRu1eDYE',
           'CAACAgIAAxkBAAENwMZnqPRdDo7mjRWCYXA9tzcIKrscAQACKwwAAiIwWEvIROJY0qdhFDYE',
           'CAACAgIAAxkBAAENwMhnqPRxIn9sDyKHfTJ5TOCgzUDi1gACUg0AAoOa6EttjhqMTmtfqzYE',
           'CAACAgIAAxkBAAENwMlnqPRxLA9epPMyFRTcizo0kSidtgACLA0AArs76EuqM6J0TSMuQTYE',
           'CAACAgIAAxkBAAENwMxnqPST74bmP7dStj_ZOcxA0uQ4cwACJg4AAvW6EEiEWDQIzqqeEzYE',
           'CAACAgIAAxkBAAENwM5nqPScHWdLjHY7WlVC0-S3kf9w2gACIgoAAu8zAAFIBThdk4Ga-zg2BA',
           'CAACAgIAAxkBAAENwNBnqPSp0k5Ow2_NrPqEGJmgVAMfUwACTA8AArE_mUtMSxoS68HKSjYE'
]

current_date = datetime.now()
name_document = f'/data/dasha.docx'
doc = Document('/data/dasha.docx')
table = doc.tables[0]
number_defect = 1
number_row = number_defect + 1
msg_id = int
edit_msg_id = int
msg_save_photo_id = int
msg_next_id = int

location = ''
description_of_defect = ''
photo = ''
probable_cause_of_occurrence = ''
recommendations = ''



class Status(StatesGroup):
    location = State()
    description_of_defect = State()
    photo = State()
    probable_cause_of_occurrence = State()
    recommendations = State()

# inline клавиатура
async def status_kb_inline() -> InlineKeyboardMarkup:
    if number_defect == len(table.rows)-2:
        key_name = new_defects
        key = 'nd'
    else:
        key_name = edit_defects
        key = 'qst'
    builder = InlineKeyboardBuilder()
    # Добавляем кнопки
    for key_id, key_data in key_name.items():
        builder.row(InlineKeyboardButton(text=key_data.get(f'{key}'), callback_data=f'{key}_{key_id}'))
    # Настраиваем размер клавиатуры
    return builder.adjust(1).as_markup()

inlkbt_edit_defect = InlineKeyboardMarkup(inline_keyboard=[
    [InlineKeyboardButton(text=f'Редактировать дефект', callback_data='edit_defect')]])

inlkbt_save_photo = InlineKeyboardMarkup(inline_keyboard=[
    [InlineKeyboardButton(text=f'Сохранить фото', callback_data='save_photo')]])

inlkbt_next_defect = InlineKeyboardMarkup(inline_keyboard=[
    [InlineKeyboardButton(text=f'⬇ Завершить редактирование ⬇', callback_data='nd_6')]])

# Reply клавиатура
async def status_kb_reply() -> ReplyKeyboardMarkup:
    builder = ReplyKeyboardBuilder()
    # Добавляем кнопки
    for menu_id, menu_data in menu.items():
        builder.add(KeyboardButton(text=menu_data.get('ans')))
    # Настраиваем размер клавиатуры
    return builder.adjust(2).as_markup(input_field_placeholder=f'Дефект №{number_defect}')


# Меняем состояние клавиатуры
async def edit_kb():
    if number_defect == len(table.rows)-2:
        key_col = new_defects
        key = 'nd'
    else:
        key = 'qst'
        key_col = edit_defects
    qst_1 = len(table.rows[number_row].cells[1].text)
    qst_2 = len(table.rows[number_row].cells[2].text)
    qst_3 = len(table.rows[number_row].cells[3].text)
    qst_4 = len(table.rows[number_row].cells[4].text)
    qst_5 = len(table.rows[number_row].cells[5].text)
    if qst_1 > 0:
        key_col[1][f'{key}'] = 'Место расположения ✅'
    else:
        key_col[1][f'{key}'] = 'Место расположения'

    if qst_2 > 0:
        key_col[2][f'{key}'] = 'Описание дефекта ✅'
    else:
        key_col[2][f'{key}'] = 'Описание дефекта'

    if qst_3 > 0:
        key_col[3][f'{key}'] = 'Фото ✅'
    else:
        key_col[3][f'{key}'] = 'Фото'

    if qst_4 > 0:
        key_col[4][f'{key}'] = 'Вероятная причина возникновения ✅'
    else:
        key_col[4][f'{key}'] = 'Вероятная причина возникновения'

    if qst_5 > 0:
        key_col[5][f'{key}'] = 'Рекомендации ✅'
    else:
        key_col[5][f'{key}'] = 'Рекомендации'

# Меняем состояние сообщения после закрытия дефекта
async def edit_text():
    global location, description_of_defect, photo, probable_cause_of_occurrence, recommendations
    qst_1 = len(table.rows[number_row].cells[1].text)
    qst_2 = len(table.rows[number_row].cells[2].text)
    qst_3 = len(table.rows[number_row].cells[3].text)
    qst_4 = len(table.rows[number_row].cells[4].text)
    qst_5 = len(table.rows[number_row].cells[5].text)
    if qst_1 > 0:
        location = '✅'
    else:
        location = '❌'
    if qst_2 > 0:
        description_of_defect = '✅'
    else:
        description_of_defect = '❌'
    if qst_3 > 0:
        photo = '✅'
    else:
        photo = '❌'
    if qst_4 > 0:
        probable_cause_of_occurrence = '✅'
    else:
        probable_cause_of_occurrence = '❌'
    if qst_5 > 0:
        recommendations = '✅'
    else:
        recommendations = '❌'

# Вынимаем все числа из текста
async def update_number_defect(find_defect):
    global number_defect
    length = len(find_defect)
    i = 0
    while i < length:
        s_int = ''
        while i < length and '0' <= find_defect[i] <= '9':
            s_int += find_defect[i]
            i += 1
        i += 1
        if s_int != '':
            number_defect = int(s_int)





# Команда /start
@dp.message(CommandStart())
async def cmd_start (massage: types.Message):
    rand_sticker = random.choice(stickers)
    await massage.answer_sticker(sticker=rand_sticker, reply_markup=await status_kb_reply())
    # await massage.answer('Добро пожаловать в бота,\nкоторый поможет тебе создать дефектную ведомость.', reply_markup=await status_kb_reply(number_defect-1))
    await edit_kb()
    table.rows[number_row].cells[0].text = str(number_defect)
    rand_priv = random.choice(privet)
    await massage.answer(rand_priv)
    msg = await massage.answer(f'Дефект №{number_defect}', reply_markup=await status_kb_inline())
    global msg_id
    msg_id = msg.message_id

# Следующий дефект
@dp.callback_query(lambda callback: callback.data=='nd_6' or callback.data=='qst_6')
async def next_defect (callback: types.CallbackQuery):
    global number_defect, number_row
    await update_number_defect(callback.message.text)
    await edit_text()
    await callback.message.edit_text(f'Дефект №{number_defect}'
                                     f'\nМесто расположения: {location}'
                                     f'\nОписание дефекта: {description_of_defect}'
                                     f'\nФото: {photo}'
                                     f'\nВероятная причина возникновения: {probable_cause_of_occurrence}'
                                     f'\nРекомендации: {recommendations}',
                                     reply_markup=inlkbt_edit_defect)
    new_row = table.add_row().cells
    number_defect = len(table.rows)-2
    number_row = number_defect + 1
    table.rows[number_row].cells[0].text = str(number_defect)
    await edit_kb()
    await callback.message.answer('________________________________', reply_markup=await status_kb_reply())
    msg_next = await callback.message.answer(f'Дефект №{number_defect}', reply_markup=await status_kb_inline(), cache_time=2)
    global msg_next_id
    msg_next_id = msg_next.message_id

# редактирование дефекта
@dp.callback_query(F.data == 'edit_defect')
async def edit_defect(callback: types.CallbackQuery, bot: Bot):
    global msg_id, edit_msg_id, number_defect, number_row
    await bot.delete_message(
        chat_id=callback.message.chat.id,
        message_id=msg_next_id)
    await edit_text()
    await bot.edit_message_text(
        chat_id=callback.message.chat.id,
        message_id=msg_id,
        text=f'Дефект №{number_defect}'
        f'\nМесто расположения: {location}'
        f'\nОписание дефекта: {description_of_defect}'
        f'\nФото: {photo}'
        f'\nВероятная причина возникновения: {probable_cause_of_occurrence}'
        f'\nРекомендации: {recommendations}',
        reply_markup=inlkbt_edit_defect)
    await update_number_defect(callback.message.text)
    number_row = number_defect + 1
    await edit_kb()
    await callback.message.edit_text(text='~Дефект отредактирован~', parse_mode='MarkdownV2', reply_markup=None)
    await callback.message.answer('________________________________', reply_markup=await status_kb_reply())
    msg = await callback.message.answer(text=f'Вы редактируете дефект {number_defect}', reply_markup=await status_kb_inline())
    msg_id = msg.message_id

# Место расположения
@dp.callback_query(lambda callback: callback.data=='nd_1' or callback.data=='qst_1')
async def col_1(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(Status.location)
    await callback.answer('Введите место расположения')
    await callback.message.edit_text(text='Место расположения:', reply_markup=None)
@dp.message(Status.location)
async def end_1(message: Message, state: FSMContext):
    await state.update_data(location=message.text)
    data = await state.get_data()
    table.rows[number_row].cells[1].text = data["location"]
    await edit_kb()
    if number_defect != len(table.rows)-2:
        msg = await message.answer(text=f'Вы редактируете дефект №{number_defect}',reply_markup=await status_kb_inline())
    else:
        msg = await message.answer(text=f'Дефект №{number_defect}',reply_markup=await status_kb_inline())
    global msg_id
    msg_id = msg.message_id
    await state.clear()

# Описание дефекта
@dp.callback_query(lambda callback: callback.data=='nd_2' or callback.data=='qst_2')
async def col_2(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(Status.description_of_defect)
    await callback.answer('Введите описание дефекта')
    await callback.message.edit_text(text='Описание дефекта:', reply_markup=None)
@dp.message(Status.description_of_defect)
async def end_2(message: Message, state: FSMContext):
    await state.update_data(description_of_defect=message.text)
    data = await state.get_data()
    table.rows[number_row].cells[2].text = data["description_of_defect"]
    await edit_kb()
    if number_defect != len(table.rows) - 2:
        msg = await message.answer(text=f'Вы редактируете дефект №{number_defect}',
                                   reply_markup=await status_kb_inline())
    else:
        msg = await message.answer(text=f'Дефект №{number_defect}', reply_markup=await status_kb_inline())
    global msg_id
    msg_id = msg.message_id
    await state.clear()

# Фото дефекта
@dp.callback_query(lambda callback: callback.data=='nd_3' or callback.data=='qst_3' or callback.data== 'save_photo')
async def col_3(callback: types.CallbackQuery, state: FSMContext, bot: Bot):
    if callback.data=='nd_3' or callback.data=='qst_3':
        await state.set_state(Status.photo)
        await callback.answer('Добавьте фото', cache_time=2)
        msg_save_photo = await callback.message.edit_text(text='Добавьте фото:',
                                                          reply_markup=inlkbt_save_photo)
        global msg_save_photo_id
        msg_save_photo_id = msg_save_photo.message_id
    elif callback.data== 'save_photo':
        await edit_kb()
        if number_defect != len(table.rows) - 2:
            msg = await callback.message.answer(text=f'Вы редактируете дефект №{number_defect}',
                                       reply_markup=await status_kb_inline())
        else:
            msg = await callback.message.answer(text=f'Дефект №{number_defect}', reply_markup=await status_kb_inline())
        global msg_id
        msg_id = msg.message_id
        await bot.delete_message(
            chat_id=callback.message.chat.id,
            message_id=msg_save_photo_id)
        await state.clear()

    @dp.message(Status.photo)
    async def end_3(message: Message, state: FSMContext):
        await state.update_data(photo=message.photo[-1])
        data = await state.get_data()
        unique_id = data['photo'].file_unique_id
        save_path = f'/data/{unique_id}.jpg'
        await message.bot.download(file=data['photo'], destination=save_path)
        cell = table.rows[number_row].cells[3]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(save_path, width=Mm(84))
        run = paragraph.add_run(f'\nФото {unique_id}.jpg')
        run.font.size = Pt(6)
        os.remove(save_path)

# Вероятная причина возникновения
@dp.callback_query(lambda callback: callback.data=='nd_4' or callback.data=='qst_4')
async def col_4(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(Status.probable_cause_of_occurrence)
    await callback.answer('Введите причину возникновения')
    await callback.message.edit_text(text='Вероятная причина возникновения:', reply_markup=None)
@dp.message(Status.probable_cause_of_occurrence)
async def end_4(message: Message, state: FSMContext):
    await state.update_data(probable_cause_of_occurrence=message.text)
    data = await state.get_data()
    table.rows[number_row].cells[4].text = data["probable_cause_of_occurrence"]
    await edit_kb()
    if number_defect != len(table.rows) - 2:
        msg = await message.answer(text=f'Вы редактируете дефект №{number_defect}',
                                   reply_markup=await status_kb_inline())
    else:
        msg = await message.answer(text=f'Дефект №{number_defect}', reply_markup=await status_kb_inline())
    global msg_id
    msg_id = msg.message_id
    await state.clear()

# Рекомендации
@dp.callback_query(lambda callback: callback.data=='nd_5' or callback.data=='qst_5')
async def col_5(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(Status.recommendations)
    await callback.answer('Введите рекомендации')
    await callback.message.edit_text(text='Рекомендации:', reply_markup=None)
@dp.message(Status.recommendations)
async def end_5(message: Message, state: FSMContext):
    await state.update_data(recommendations=message.text)
    data = await state.get_data()
    table.rows[number_row].cells[5].text = data["recommendations"]
    await edit_kb()
    if number_defect != len(table.rows) - 2:
        msg = await message.answer(text=f'Вы редактируете дефект №{number_defect}',
                                   reply_markup=await status_kb_inline())
    else:
        msg = await message.answer(text=f'Дефект №{number_defect}', reply_markup=await status_kb_inline())
    global msg_id
    msg_id = msg.message_id
    await state.clear()

# Сохранение файла
@dp.message(F.text == 'Сохранить файл 📝')
async def save_file (message: Message):
    doc.save(name_document)
    await message.answer('Файл сохранен')

# Загрузка файла
@dp.message(F.text == 'Загрузить файл 📖')
async def download_file (message: types.Message):
    document =  FSInputFile(name_document)
    await message.answer_document (document=document, caption='Ты молодчина 💪')
    os.remove(name_document)

# Создать новый файл
@dp.message(F.text == 'Создать новый проект 🆕')
async def new_file (message: Message):
    await message.answer('Новый файл создан')
    os.system('python3 run.py')











async def main():

    await dp.start_polling(bot)



if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO)
    try:
        asyncio.run(main())
    except KeyboardInterrupt:

        print('Бот выключен')
